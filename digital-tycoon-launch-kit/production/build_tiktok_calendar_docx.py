from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE = Path('/home/ubuntu/tycoon_launch_assets/14_Day_TikTok_Traffic_Calendar.md')
OUTPUT = Path('/home/ubuntu/tycoon_launch_assets/14_Day_TikTok_Traffic_Calendar.docx')

BRASS = 'B9852E'
INK = '16130E'
PAPER = 'FBF6EC'
SOFT = 'E9DDC7'
CHARCOAL = '25211A'


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def add_text_with_bold(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_table(document, rows):
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = 'Table Grid'
    table.autofit = True
    for r_index, values in enumerate(rows):
        cells = table.add_row().cells
        for c_index, value in enumerate(values):
            cell = cells[c_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell)
            shade_cell(cell, CHARCOAL if r_index == 0 else (PAPER if r_index % 2 else SOFT))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            add_text_with_bold(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(8.7 if r_index else 8.3)
                run.font.name = 'Aptos'
                run.font.color.rgb = RGBColor.from_string(PAPER if r_index == 0 else INK)
                if r_index == 0:
                    run.bold = True
        if r_index == 0:
            set_repeat_table_header(table.rows[0])
    document.add_paragraph()


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line or re.fullmatch(r'\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)*\|?', line):
            continue
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        rows.append(cells)
    return rows


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    normal = doc.styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 24), (2, 17), (3, 13)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Georgia'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(15 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(6)

    quote = doc.styles['Quote']
    quote.font.name = 'Aptos'
    quote.font.size = Pt(9.5)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string('5A4A34')
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.space_after = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('THE DIGITAL TYCOON PLAYBOOK  •  14-DAY TIKTOK TRAFFIC ENGINE')
    run.font.name = 'Aptos'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string('73531B')


def build_document():
    doc = Document()
    configure_document(doc)

    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            rows = parse_table(table_lines)
            if rows and all(len(row) == len(rows[0]) for row in rows):
                add_table(doc, rows)
            table_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith('|'):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip() or line.strip() == '---':
            continue
        if line.startswith('# '):
            title = doc.add_heading(line[2:].strip(), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in title.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(27)
                run.font.color.rgb = RGBColor.from_string(INK)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            add_text_with_bold(p, line[2:].strip())
            continue
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_bold(p, line[2:].strip())
            continue

        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        label = re.match(r'^(\*\*[^*]+\*\*:\s*)(.*)$', line)
        if label:
            add_text_with_bold(p, label.group(1) + label.group(2))
        else:
            add_text_with_bold(p, line)

    flush_table()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    build_document()
